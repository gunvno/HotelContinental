from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "HotelContinental_API_NghiepVu_ChucNang.docx"


ACCENT = "A8651E"
TEXT = "0B1B35"
MUTED = "6B5B4B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], ACCENT)
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor.from_string(TEXT if level == 1 else ACCENT)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10)


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10)


def parse_endpoints() -> list[dict[str, str]]:
    endpoints: list[dict[str, str]] = []
    class_re = re.compile(r"@RequestMapping\((?:value\s*=\s*)?[\"']([^\"']+)[\"']\)")
    inline_re = re.compile(r"@(GetMapping|PostMapping|PutMapping|PatchMapping|DeleteMapping|RequestMapping)(?:\((.*?)\))?")
    path_re = re.compile(r"(?:value\s*=\s*)?(?:\{\s*)?[\"']([^\"']*)[\"']")
    method_re = re.compile(r"method\s*=\s*RequestMethod\.([A-Z]+)")
    method_map = {
        "GetMapping": "GET",
        "PostMapping": "POST",
        "PutMapping": "PUT",
        "PatchMapping": "PATCH",
        "DeleteMapping": "DELETE",
        "RequestMapping": "ANY",
    }

    for path in sorted((ROOT / "backend").glob("*-service/src/main/java/**/*Controller.java")):
        text = path.read_text(encoding="utf-8", errors="ignore")
        service = path.parts[path.parts.index("backend") + 1]
        class_base = ""
        m_base = class_re.search(text)
        if m_base:
            class_base = m_base.group(1)

        pending: list[tuple[str, str]] = []
        for raw in text.splitlines():
            line = raw.strip()
            m = inline_re.search(line)
            if m:
                ann, args = m.group(1), m.group(2) or ""
                method = method_map[ann]
                method_match = method_re.search(args)
                if method_match:
                    method = method_match.group(1)
                found_path = ""
                path_match = path_re.search(args)
                if path_match:
                    found_path = path_match.group(1)
                if not (ann == "RequestMapping" and line.startswith("@RequestMapping")):
                    pending.append((method, found_path))
                continue

            if pending and re.search(r"\b(public|private|protected)?\s*[^;=]+\([^;]*\)\s*\{?", line):
                name = line.split("(")[0].split()[-1]
                for method, suffix in pending:
                    full = "/" + "/".join([class_base.strip("/"), suffix.strip("/")]).strip("/")
                    full = full.replace("//", "/")
                    endpoints.append(
                        {
                            "service": service,
                            "method": method,
                            "endpoint": full,
                            "handler": name,
                        }
                    )
                pending = []

    priority = {
        "identity-service": 0,
        "booking-service": 1,
        "room-service": 2,
        "billing-service": 3,
        "catalog-service": 4,
        "promotion-service": 5,
        "report-service": 6,
        "chat-service": 7,
        "content-service": 8,
        "feedback-service": 9,
        "ai-assistant-service": 10,
    }
    return sorted(endpoints, key=lambda e: (priority.get(e["service"], 99), e["endpoint"], e["method"]))


def describe_endpoint(ep: dict[str, str]) -> str:
    service, method, path = ep["service"], ep["method"], ep["endpoint"]
    if service == "identity-service":
        if path.startswith("/auth/"):
            return "Đăng nhập, đăng ký, refresh token, logout hoặc lấy thông tin người dùng."
        if path.startswith("/admin/permissions"):
            return "Quản lý tài khoản nhân viên, role, trạng thái và mật khẩu."
        if path.startswith("/staff-activity"):
            return "Theo dõi đăng nhập, check-in ca, check-out ca của nhân viên."
        return "Tra cứu thông tin người dùng."
    if service == "booking-service":
        if "change-dates" in path:
            return "Đổi ngày lưu trú và ghi lịch sử chỉnh sửa booking."
        if "check-in" in path or "check-out" in path:
            return "Xử lý nhận phòng hoặc trả phòng."
        if "cancel" in path:
            return "Khách yêu cầu hủy hoặc admin duyệt hủy booking."
        if "edit-history" in path:
            return "Xem lịch sử thay đổi booking."
        return "Tạo, xem và vận hành booking phòng."
    if service == "room-service":
        if "housekeeping" in path:
            return "Quản lý trạng thái phòng và phân công dọn phòng."
        if path.startswith("/building"):
            return "Quản lý tòa nhà, tầng và sơ đồ phòng."
        return "Quản lý phòng, ảnh phòng và trạng thái kinh doanh."
    if service == "billing-service":
        if "service-order-details" in path:
            return "Quản lý dịch vụ phát sinh, phân công, duyệt và phục vụ."
        if "payment" in path:
            return "Tạo yêu cầu thanh toán, xử lý PayOS hoặc ghi nhận thanh toán."
        if "invoice" in path:
            return "Lấy hóa đơn theo booking."
        return "Tổng hợp chi phí và thanh toán."
    if service == "catalog-service":
        return "Quản lý danh mục gốc: loại phòng, tiện nghi, dịch vụ, bảng giá và dịch vụ theo loại phòng."
    if service == "promotion-service":
        return "Quản lý voucher và áp dụng ưu đãi."
    if service == "report-service":
        return "Tổng hợp doanh thu và chỉ số vận hành."
    if service == "chat-service":
        return "Quản lý hội thoại hỗ trợ khách hàng."
    if service == "content-service":
        return "Quản lý nội dung chính sách khách sạn."
    if service == "feedback-service":
        return "Ghi nhận và xem đánh giá của khách."
    if service == "ai-assistant-service":
        return "Hội thoại trợ lý AI cho khách."
    return "Endpoint nghiệp vụ."


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HOTEL CONTINENTAL")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tài liệu API, luồng nghiệp vụ và chức năng hệ thống")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phiên bản nội bộ - mô hình ERP một khách sạn")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Cập nhật: {date.today().strftime('%d/%m/%Y')}")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)

    add_cover(doc)

    add_heading(doc, "1. Tổng quan hệ thống", 1)
    add_bullets(
        doc,
        [
            "Hotel Continental được định hướng thành hệ thống quản trị một khách sạn, không còn bài toán nhiều khách sạn.",
            "Website khách hàng vẫn giữ vai trò đặt phòng, thanh toán booking ban đầu, xem hóa đơn và gọi thêm dịch vụ trong thời gian lưu trú.",
            "Admin house là cổng vận hành nội bộ theo hướng ERP: đặt phòng, lễ tân, buồng phòng, dịch vụ phát sinh, nhân viên, doanh thu và báo cáo.",
            "Backend tách thành nhiều service Spring Boot, giao tiếp qua API Gateway và dần chuyển các luồng liên service sang Kafka event.",
            "Dữ liệu lõi xoay quanh phòng, booking, hóa đơn, thanh toán, dịch vụ phát sinh, phân công việc và lịch sử hoạt động nhân viên.",
        ],
    )

    add_heading(doc, "2. Vai trò và phân quyền", 1)
    add_table(
        doc,
        ["Role", "Mục đích", "Quyền chính"],
        [
            ["ADMIN", "Quản trị hệ thống", "Tạo tài khoản nhân viên, gán role, cấu hình danh mục, xem toàn bộ dữ liệu."],
            ["MANAGER", "Quản lý vận hành", "Theo dõi dashboard, duyệt hủy, xem báo cáo, điều phối lễ tân, dịch vụ và buồng phòng."],
            ["RECEPTIONIST", "Lễ tân", "Check-in, check-out, thêm dịch vụ vào booking, xử lý thanh toán còn phải thu."],
            ["CUSTOMER_SUPPORT", "Chăm sóc khách hàng", "Quản lý hội thoại, hỗ trợ khách, theo dõi yêu cầu hủy hoặc đổi ngày."],
            ["HOUSEKEEPING", "Buồng phòng", "Xem việc được giao, nhận việc, hoàn thành dọn phòng hoặc phục vụ dịch vụ."],
            ["CUSTOMER", "Khách hàng", "Đặt phòng, thanh toán, xem hóa đơn, đổi ngày nếu đủ điều kiện, yêu cầu hủy, gọi thêm dịch vụ."],
        ],
        [3, 5, 10],
    )

    add_heading(doc, "3. Mô tả nghiệp vụ", 1)
    add_bullets(
        doc,
        [
            "Booking ban đầu: khách chọn phòng theo ngày hoặc theo giờ, hệ thống kiểm tra phòng trống theo khung thời gian, sau đó khách thanh toán toàn bộ tiền phòng ban đầu.",
            "Thanh toán: booking được giữ ở trạng thái chờ thanh toán đến khi PayOS/webhook hoặc mock-paid xác nhận thành công; sau đó booking chuyển sang đã đặt/đã giữ phòng.",
            "Đổi ngày: chỉ cho đổi khi còn trước giờ nhận phòng theo chính sách và phải giữ nguyên số ngày/số giờ đã thanh toán.",
            "Hủy booking: khách gửi yêu cầu hủy; admin hoặc manager duyệt theo chính sách hoàn tiền/không hoàn tiền.",
            "Check-in/check-out: lễ tân xác nhận khách nhận phòng, đăng ký lưu trú, sau đó khi checkout hệ thống tổng hợp tiền dịch vụ phát sinh, phụ thu và số còn phải trả.",
            "Dịch vụ phát sinh: khách hoặc lễ tân tạo yêu cầu dịch vụ; dịch vụ có thể thanh toán ngay bằng PayOS hoặc ghi vào hóa đơn checkout.",
            "Phân công vận hành: housekeeping và service order có người nhận việc, người giao, thời điểm nhận, thời điểm hoàn thành để quản lý công việc thật.",
            "Audit/history: thay đổi quan trọng của booking như đổi ngày, cập nhật tiền, hủy, check-in/check-out cần được ghi vào lịch sử chỉnh sửa.",
        ],
    )

    add_heading(doc, "4. Luồng nghiệp vụ chính", 1)
    flows = [
        (
            "4.1. Khách đặt phòng và thanh toán ban đầu",
            [
                "Khách vào website, chọn chế độ theo đêm hoặc theo giờ.",
                "Hệ thống mặc định ngày nhận phòng là hôm nay và kiểm tra phòng trống theo khung thời gian.",
                "Khách chọn phòng, nhập thông tin, áp voucher nếu có.",
                "Booking được tạo với trạng thái chờ thanh toán.",
                "Billing service tạo payment request PayOS và trả QR/thông tin chuyển khoản.",
                "Khi thanh toán thành công, booking được đánh dấu đã thanh toán/đã giữ phòng và hóa đơn được cập nhật.",
            ],
        ),
        (
            "4.2. Khách đổi ngày lưu trú",
            [
                "Khách mở chi tiết booking trong trang tài khoản.",
                "Hệ thống kiểm tra chính sách đổi ngày miễn phí và trạng thái phòng.",
                "Nếu booking theo ngày thì chỉ cho đổi ngày; nếu booking theo giờ mới cho đổi giờ.",
                "Số ngày hoặc số giờ mới phải bằng thời lượng đã thanh toán.",
                "Booking service cập nhật ngày, ghi edit_history và phát event cho service liên quan.",
            ],
        ),
        (
            "4.3. Check-in và check-out",
            [
                "Lễ tân mở chi tiết booking trong admin.",
                "Khi khách đến, lễ tân bấm check-in và có thể đăng ký khách lưu trú.",
                "Phòng chuyển sang đang ở; dashboard ghi nhận phòng đang ở và sắp check-out.",
                "Khi checkout, hệ thống tổng hợp tiền phòng, dịch vụ phát sinh, phụ thu, voucher, đặt cọc và số còn phải trả.",
                "Lễ tân chọn phương thức thanh toán phần còn lại, xác nhận check-out và phòng chuyển sang cần dọn.",
            ],
        ),
        (
            "4.4. Khách gọi thêm dịch vụ",
            [
                "Khách vào Tài khoản > Gọi thêm dịch vụ.",
                "Khách chọn booking đang lưu trú, chọn dịch vụ, số lượng, ghi chú và chọn cách thanh toán.",
                "Nếu thanh toán PayOS ngay, hệ thống tạo payment request riêng cho service order.",
                "Nếu ghi vào phòng, khoản này được cộng vào hóa đơn checkout.",
                "Admin nhìn thấy service order trong trang quản lý dịch vụ phát sinh để duyệt, phân công và đánh dấu đã phục vụ.",
            ],
        ),
        (
            "4.5. Housekeeping nhận và hoàn thành việc",
            [
                "Sau checkout hoặc khi quản lý tạo việc, phòng có trạng thái cần dọn.",
                "Manager hoặc receptionist phân công nhân viên housekeeping cho phòng.",
                "Housekeeping mở trang việc của mình, nhận thông tin phòng, người giao và thời điểm giao.",
                "Khi hoàn thành, housekeeping bấm hoàn thành; hệ thống ghi người hoàn thành, thời điểm hoàn thành và cập nhật trạng thái phòng.",
            ],
        ),
        (
            "4.6. Theo dõi hoạt động nhân viên",
            [
                "Nhân viên đăng nhập vào hệ thống bằng tài khoản được admin tạo.",
                "Nhân viên check-in ca làm và check-out ca làm trong hệ thống.",
                "Identity service ghi lại thời điểm đăng nhập, check-in, check-out và trạng thái phiên làm việc.",
                "Manager dùng dữ liệu này để xem lịch sử làm việc, tính công sơ bộ và kiểm tra trách nhiệm khi xử lý nghiệp vụ.",
            ],
        ),
    ]
    for title, steps in flows:
        add_heading(doc, title, 2)
        add_steps(doc, steps)

    add_heading(doc, "5. Chức năng app khách hàng", 1)
    add_table(
        doc,
        ["Nhóm chức năng", "Mô tả"],
        [
            ["Trang chủ", "Giới thiệu khách sạn, điều hướng đến danh sách phòng, tiện nghi và liên hệ."],
            ["Phòng & Suite", "Tìm phòng theo ngày/giờ, số khách, khoảng giá; xem chi tiết phòng và đặt phòng."],
            ["Thanh toán booking", "Hiển thị QR PayOS, thông tin chuyển khoản, trạng thái chờ thanh toán và màn hình thành công."],
            ["Hồ sơ tài khoản", "Xem thông tin cá nhân, lịch sử hóa đơn, booking đã đặt và trạng thái lưu trú."],
            ["Gọi thêm dịch vụ", "Khách chọn booking đang ở, chọn dịch vụ, thanh toán ngay hoặc ghi vào phòng."],
            ["Hỗ trợ", "Nhắn tin với chăm sóc khách hàng hoặc trợ lý AI nếu bật."],
        ],
        [5, 13],
    )

    add_heading(doc, "6. Chức năng app admin", 1)
    add_table(
        doc,
        ["Màn hình", "Người dùng chính", "Chức năng"],
        [
            ["Dashboard vận hành", "Admin, Manager", "Phòng đang ở, sắp check-in/check-out, booking chờ hủy, doanh thu hôm nay/tháng."],
            ["Đặt phòng", "Receptionist, Manager", "Danh sách booking, chi tiết booking, check-in, check-out, hủy, duyệt hủy, lịch sử chỉnh sửa."],
            ["Phòng", "Manager, Receptionist", "Danh sách phòng, trạng thái kinh doanh, housekeeping status, ảnh phòng."],
            ["Housekeeping", "Manager, Housekeeping", "Lọc việc dọn phòng, phân công, xem người nhận, hoàn thành việc."],
            ["Dịch vụ phát sinh", "Receptionist, Manager, Housekeeping", "Tạo dịch vụ cho booking, duyệt, phân công, phục vụ, lọc theo trạng thái."],
            ["Dịch vụ gốc", "Admin, Manager", "Thêm/sửa/xóa dịch vụ, đổi giá, bật/tắt trạng thái bán."],
            ["Gán dịch vụ theo loại phòng", "Admin, Manager", "Gán dịch vụ kèm theo loại phòng và số lượng mặc định."],
            ["Nhân viên & quyền", "Admin", "Tạo tài khoản, role, trạng thái, reset mật khẩu và quyền nhân viên."],
            ["Hoạt động nhân viên", "Admin, Manager", "Xem lịch sử đăng nhập, check-in ca, check-out ca."],
            ["Doanh thu/Báo cáo", "Admin, Manager", "Tổng hợp doanh thu theo ngày/tháng và khoảng thời gian."],
            ["Tin nhắn", "Customer Support", "Danh sách hội thoại, trả lời, đánh dấu đã đọc, đóng hội thoại."],
        ],
        [4, 4, 10],
    )

    add_heading(doc, "7. API tổng hợp", 1)
    p = doc.add_paragraph()
    p.add_run("Quy ước gọi qua gateway: ").bold = True
    p.add_run("http://localhost:8888/api/v1/{service-context}{endpoint}. Ví dụ booking service dùng /api/v1/booking/room-bookings.")

    endpoints = parse_endpoints()
    rows = []
    for ep in endpoints:
        if ep["endpoint"] == "/health":
            continue
        rows.append(
            [
                ep["service"].replace("-service", ""),
                ep["method"],
                ep["endpoint"],
                describe_endpoint(ep),
            ]
        )
    add_table(doc, ["Service", "Method", "Endpoint", "Mục đích"], rows, [2.7, 1.6, 5.2, 8.2])

    add_heading(doc, "8. Trạng thái và dữ liệu quan trọng", 1)
    add_table(
        doc,
        ["Nhóm", "Trạng thái/dữ liệu", "Ý nghĩa"],
        [
            ["Booking", "PENDING_PAYMENT, BOOKED/DEPOSITED, CHECKED_IN, CHECKED_OUT, CANCEL_REQUESTED, CANCELED", "Theo dõi vòng đời đặt phòng từ lúc tạo đến khi hoàn tất hoặc hủy."],
            ["Room", "AVAILABLE, OCCUPIED, DIRTY, MAINTENANCE", "Quản lý khả năng bán phòng và trạng thái buồng phòng."],
            ["Housekeeping", "assignedTo, assignedBy, assignedTime, completedBy, completedTime", "Ghi nhận ai nhận việc, ai giao việc và khi nào hoàn thành."],
            ["Service order", "PENDING, APPROVED, REJECTED, SERVED, PAID_AT_CHECKOUT", "Theo dõi yêu cầu dịch vụ phát sinh từ lúc tạo đến khi phục vụ/thanh toán."],
            ["Payment", "PENDING, PAID, FAILED/EXPIRED, REFUND_PENDING/REFUNDED", "Theo dõi thanh toán PayOS, chuyển khoản và hoàn tiền nếu có."],
            ["Audit", "edit_history", "Ghi lại trường bị sửa, giá trị cũ, giá trị mới, người sửa và thời điểm sửa."],
        ],
        [4, 6, 9],
    )

    add_heading(doc, "9. Event/Kafka", 1)
    add_bullets(
        doc,
        [
            "Booking service nên phát event khi tạo booking, thanh toán thành công, đổi ngày, check-in, check-out và hủy.",
            "Billing service consume booking event để tạo/cập nhật invoice, payment snapshot và tổng tiền.",
            "Room service consume check-in/check-out event để cập nhật trạng thái phòng và housekeeping.",
            "Notification service consume event để gửi email/thông báo cho khách hoặc nhân viên.",
            "Report service consume payment/booking event để tạo snapshot doanh thu, tránh phải query chéo service trực tiếp.",
        ],
    )

    add_heading(doc, "10. Ghi chú triển khai", 1)
    add_bullets(
        doc,
        [
            "MySQL đang dùng local; khi đổi entity hoặc enum cần kiểm tra lại schema hiện tại, đặc biệt các cột enum/string như payment_method.",
            "Docker compose backend hiện phục vụ mục tiêu khởi động nhanh service phụ trợ như Kafka, Zookeeper, Redis và các service backend nếu máy đủ tài nguyên.",
            "Frontend admin và customer có thể chạy riêng bằng npm để phát triển giao diện nhanh hơn.",
            "Khi thêm role hoặc quyền mới cần seed lại identity_service hoặc cập nhật dữ liệu quyền hiện có.",
            "Các thông báo UI nên dùng toast góc trên bên phải, tự ẩn sau 5 giây để đồng bộ trải nghiệm.",
        ],
    )

    add_heading(doc, "11. Đề xuất phát triển tiếp", 1)
    add_table(
        doc,
        ["Ưu tiên", "Hạng mục", "Lý do"],
        [
            ["1", "Chuẩn hóa OpenAPI/Swagger cho từng service", "Tự sinh tài liệu API chính xác hơn và hỗ trợ test endpoint."],
            ["2", "Hoàn thiện audit log toàn hệ thống", "Biết ai sửa booking, hóa đơn, dịch vụ, phòng và quyền nhân viên."],
            ["3", "Bảng ca làm việc và chấm công", "Biến staff activity thành module nhân sự ERP rõ ràng hơn."],
            ["4", "Kho vật tư buồng phòng/dịch vụ", "Theo dõi tồn kho minibar, giặt là, đồ amenities, chi phí vận hành."],
            ["5", "Luồng hoàn tiền và đối soát PayOS", "Đóng nghiệp vụ tài chính khi hủy booking hoặc phát sinh hoàn tiền."],
            ["6", "Notification thật qua email/SMS", "Không chỉ log consumer mà gửi thông báo cho khách và nhân viên."],
        ],
        [2, 6, 10],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
